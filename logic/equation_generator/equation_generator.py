import os

import matplotlib.pyplot as plt
from docx import Document


from docx.shared import Inches
from logic.equation.linear_equation import LinearEquation, Difficult


class EquationGenerator:
    def __init__(self, difficult: Difficult, quantity_of_equations: int = 10):
        self.quantity_of_equations = quantity_of_equations
        self.difficult = difficult

    def generate_equations(self) -> list[LinearEquation]:
        equations = [
            LinearEquation.from_template_with_concrete_difficult(self.difficult)
            for _ in range(self.quantity_of_equations)
        ]
        return equations

    @staticmethod
    def delete_all_pngs():
        current_directory = os.getcwd()
        files = os.listdir(current_directory)
        for file in files:
            if file.endswith(".png"):
                os.remove(os.path.join(current_directory, file))
                print(f"Удален файл: {file}")

    @staticmethod
    def __create_equation_image(
        equation: str, filename: str = "equation_image.png"
    ) -> None:
        fig, ax = plt.subplots(figsize=(2, 1))
        ax.text(0.5, 0.5, f"{equation}", fontsize=20, ha="center", va="center")
        ax.axis("off")
        plt.savefig(filename, bbox_inches="tight")
        plt.close()

    @staticmethod
    def __convert_python_string_to_latex(equation: str) -> str:
        latex_equation = equation.replace("*", r"\cdot ")
        latex_equation = f"${latex_equation}$"
        return latex_equation

    def __add_equation_with_number(self, doc, equation, number):
        latex_equation = self.__convert_python_string_to_latex(str(equation))
        image_filename = f"equation_{number}.png"
        self.__create_equation_image(latex_equation, image_filename)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        row = table.rows[0].cells
        row[0].text = f"{number}"
        paragraph = row[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_filename, width=Inches(2))

    def generate_equations_to_docx_file(
        self, equations: list[LinearEquation], filename: str = "equation.doc"
    ) -> None:
        doc = Document()
        doc.add_heading(
            f"Уравнения с уровнем сложности {self.difficult.value}", level=1
        )

        for i, equation in enumerate(equations, 1):
            self.__add_equation_with_number(doc, equation.equation, i)
            doc.add_paragraph()
        doc.save(filename)
        self.delete_all_pngs()

    def generate_answers_to_docx_file(
        self, equations: list[LinearEquation], filename: str = "answers.doc"
    ):
        doc = Document()
        doc.add_heading("Ответы для уравнений", level=1)

        for i, equation in enumerate(equations, 1):
            self.__add_equation_with_number(doc, equation.answer, i)
            doc.add_paragraph()
        doc.save(filename)
        self.delete_all_pngs()


if __name__ == "__main__":
    g = EquationGenerator(difficult=Difficult.EASY, quantity_of_equations=20)
    eq = g.generate_equations()
    g.generate_equations_to_docx_file(eq)
    g.generate_answers_to_docx_file(eq)
